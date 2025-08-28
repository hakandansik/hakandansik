##### VERTICAL RESERVATION SYSTEM V3.1 #####
from flask import current_app
from flask import Flask, render_template, request, redirect, flash, jsonify, session, url_for, send_file, send_from_directory
from apscheduler.schedulers.background import BackgroundScheduler
from pymongo import MongoClient
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, date
import os
from bson.objectid import ObjectId
import pandas as pd
from io import BytesIO
from flask import make_response
import pytz
import hashlib
import base64
from werkzeug.security import check_password_hash
from dotenv import load_dotenv
import secrets
from docx import Document
import re
import tempfile
import time
import unicodedata
import pdfkit
from flask_wtf import CSRFProtect


app = Flask(__name__, static_folder='static')
csrf = CSRFProtect(app)
# Rastgele 24 byte uzunluğunda bir secret key oluşturur
app.secret_key = os.urandom(24)
# Ortam değişkeninden SECRET_KEY sağlanmışsa onu kullan
if os.getenv("SECRET_KEY"):
    app.secret_key = os.getenv("SECRET_KEY")


# MongoDB Bağlantısı
client = MongoClient("mongodb://127.0.0.1:27017")
db = client["local"]
customer_collection = db["Rezervasyon"]
users_collection = db["users"]
kvkk_pdf_koleksiyon = db["kvkk_pdf"]
tur_collection = db["turlar"]

# wkhtmltopdf.exe dosyasının tam yolunu belirtin.
# Kurulum yaptığınız yolu buraya yazın.
path_wkhtmltopdf = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)
TEMPLATE_PATH = os.path.join(app.root_path, "static", "sahil_guvenlik.docx")


def _min_year_from_db():
    # Rezervasyon
    p1 = customer_collection.aggregate([
        {"$group": {"_id": None, "minDate": {"$min": "$TARIH_NESNESI"}}}
    ])
    min1 = next(p1, {}).get("minDate")

    # Geçmis Rezervasyon
    p2 = db["gecmis_rezervasyon"].aggregate([
        {"$group": {"_id": None, "minDate": {"$min": "$TARIH_NESNESI"}}}
    ])
    min2 = next(p2, {}).get("minDate")

    candidates = [d for d in [min1, min2] if d]
    if not candidates:
        return 2023  # emniyetli varsayilan alt sinir (istersen degistir)

    return min(candidates).year


def _norm_tr(s: str) -> str:
    """
    Türkçe karakter/aksan normalize + güvenli kıyaslama için sadeleştirir.
    Örn: 'Dalış Tarihi:' -> 'dalis tarihi:'
    """
    s = "" if s is None else str(s)
    s = s.casefold()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch)
                != "Mn")  # aksanları kaldır
    s = s.replace("ı", "i").replace("İ", "i")
    s = re.sub(r'[^a-z0-9./:()\s-]', " ", s)   # sadeleştir
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _find_table_by_headers(doc: Document, headers):
    """Tabloları gezip ilk satırda verilen tüm başlıkları içeren tabloyu bulur."""
    want = [_norm_tr(h) for h in headers]
    for tbl in doc.tables:
        if len(tbl.rows) == 0:
            continue
        # Başlıkların sağlam bir şekilde karşılaştırılması için _norm_tr kullan
        row0_texts = [_norm_tr(c.text) for c in tbl.rows[0].cells]
        if all(any(h in cell for cell in row0_texts) for h in want):
            return tbl
    return None


def _append_after_colon(paragraph, value: str):
    """
    Paragraf metninin sonuna bir run ekler (örn: '...: ' + value).
    Mevcut biçimlendirme bozulmaz.
    """
    if not value:
        return
    # Sonda ':' yoksa yine de boşlukla ekleyelim (dokümanın mevcut hali ': ' içeriyor)
    run_text = (" " if paragraph.text.endswith(":") else ": ") + str(value)
    paragraph.add_run(run_text)


def _prepare_sg_form_data(rez: dict) -> dict:
    """
    Sahil Güvenlik formları (önizleme ve DOCX) için gereken verileri hazırlar.
    """
    tur_bilgisi = rez.get("TUR_BILGISI", "")
    yer, tarih_norm, saat_str = parse_tur_bilgisi(tur_bilgisi)

    program_tarih = tarih_norm or rez.get("TARIH", "") or ""
    program_yeri = (yer or tur_bilgisi or "").strip()
    program_saat = saat_str or "10:00"

    # Yaş için referans tarih
    if isinstance(rez.get("TARIH_NESNESI"), datetime):
        ref_date = rez["TARIH_NESNESI"].date()
    else:
        ref_date = _parse_date_multiformat(program_tarih) or date.today()

    milliyet_default = (rez.get("MILLIYET") or "T.C.").strip()

    # Dinamik dalgıç listesi
    participants = []
    dalgiclar_raw = rez.get("DALGICLAR")
    if isinstance(dalgiclar_raw, list) and dalgiclar_raw:
        for d in dalgiclar_raw:
            participants.append({
                "ad": (d.get("AD") or "").strip(),
                "belgesi": (d.get("SERTIFIKA") or "").strip(),
                "tc": (d.get("TC_PASAPORT") or "").strip(),
                "milliyet": (d.get("MILLIYET") or milliyet_default).strip(),
                "yas": calculate_age(d.get("DOGUM_TARIHI", ""), ref_date)
            })
    else:
        # Tek kişi rezervasyon
        participants.append({
            "ad": (rez.get("AD") or "").strip(),
            "belgesi": (rez.get("SERTIFIKA") or "").strip(),
            "tc": (rez.get("TC_PASAPORT") or "").strip(),
            "milliyet": milliyet_default,
            "yas": calculate_age(rez.get("DOGUM_TARIHI", ""), ref_date)
        })

    # Sablondaki ilk 2 satır sabit olduğundan, aynı isimler gelirse tekrar etmeyelim
    SABIT_ISIMLER = {"ERCÜMENT SEN", "BILGE SEN", "ERCÜMENT ŞEN", "BİLGE ŞEN"}
    participants = [
        p for p in participants if p.get("ad", "").strip().upper() not in SABIT_ISIMLER
    ]

    return {
        "program_yeri": program_yeri,
        "program_tarih": program_tarih,
        "program_saat": program_saat,
        "participants": participants,
    }


def _get_rezervasyonlar(match_filter: dict, projection: dict | None = None) -> list:
    """
    Verilen filtrelere göre rezervasyonları ve kredi bilgilerini MongoDB'den çeker.
    """
    pipeline = [
        {
            '$match': match_filter
        },
        {
            '$lookup': {
                'from': 'kredi',
                'localField': 'TC_PASAPORT',
                'foreignField': 'TC_PASAPORT',
                'as': 'kredi_bilgisi_arr'
            }
        },
        {
            '$unwind': {
                'path': '$kredi_bilgisi_arr',
                'preserveNullAndEmptyArrays': True
            }
        },
        {
            '$sort': {
                'TARIH_NESNESI': 1
            }
        }
    ]
    if projection:
        pipeline.append({'$project': projection})
    return list(customer_collection.aggregate(pipeline))


def _create_and_save_rezervasyon(form_data: dict, tur: dict, kredi_kullanildi: bool, kalan_kredi_sonrasi: int) -> dict:
    """
    Rezervasyon belgesini oluşturur, veritabanına kaydeder ve oluşturulan belgeyi döndürür.
    """
    # Tur başlangıç/bitiş tarihleri
    tur_baslangic_dt = tur.get("BASLANGIC_TARIHI")
    tur_bitis_dt = tur.get("BITIS_TARIHI")
    tarih_sql_format = tur_baslangic_dt.strftime('%d/%m/%Y')
    bitis_sql_format = tur_bitis_dt.strftime(
        '%d/%m/%Y') if tur_bitis_dt else tarih_sql_format

    # Tur bilgisi
    tur_isim = (tur.get("TUR_ISMI") or "").strip()
    tur_bilgisi = f"{tur_isim} {tarih_sql_format}".strip()

    # KVKK bilgileri
    kvkk_doc = kvkk_pdf_koleksiyon.find_one(sort=[("VERSIYON", -1)])
    kvkk_versiyon = kvkk_doc["VERSIYON"] if kvkk_doc else "Bilinmiyor"
    pdf_path = os.path.join(app.static_folder,
                            "ACIK RIZA METNI_VERTICAL.pdf")
    hash_degeri = hesapla_pdf_sha256(
        pdf_path) if os.path.exists(pdf_path) else None

    rezervasyon = {
        "AD": form_data.get('adsoyad'),
        "GSM": form_data.get('telefon'),
        "EMAIL": form_data.get('email'),
        "TC_PASAPORT": form_data.get('tc_pasaport'),
        "DOGUM_TARIHI": form_data.get('dogumtarihi') or form_data.get('dogum_tarihi'),
        "ADRES": form_data.get('adres'),
        "SERTIFIKA": form_data.get('sertifika'),
        "TARIH": tarih_sql_format,
        "TARIH_NESNESI": tur_baslangic_dt,
        "REZ_BITIS": bitis_sql_format,
        "REZ_BITIS_NESNESI": tur_bitis_dt or tur_baslangic_dt,
        "IPTAL_TOKEN": secrets.token_urlsafe(24),
        "KAYIT_TARIHI": datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
        "TUR_ID": tur["_id"],
        "TUR_BILGISI": tur_bilgisi,
        "KREDI": kalan_kredi_sonrasi if kredi_kullanildi else "Kredi Yok",
        "REZERVASYON_DURUMU": "AKTIF",
        "OLUSTURMA_TARIHI": datetime.now(),
        "KVKK": {
            "ONAY": True,
            "ONAY_TARIHI": form_data.get('kvkk_onay_tarihi') or datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
            "IP": request.remote_addr,
            "VERSIYON": kvkk_versiyon,
            "HASH": hash_degeri
        }
    }
    customer_collection.insert_one(rezervasyon)
    return rezervasyon


def _send_rezervasyon_emails(rezervasyon: dict, tur: dict):
    """
    Rezervasyon sonrası yönetici ve kullanıcıya bilgilendirme e-postalarını gönderir.
    """
    try:
        kapasite = tur.get("KONTENJAN", 0)
        yeni_kayitli_sayisi = customer_collection.count_documents(
            {"TUR_ID": tur["_id"]})
        kalan_kontenjan = max(0, kapasite - yeni_kayitli_sayisi)

        eposta_bilgileri = {
            "adsoyad": rezervasyon["AD"], "telefon": rezervasyon["GSM"],
            "email": rezervasyon["EMAIL"], "tarih": rezervasyon["TARIH"],
            "tur_bilgisi": rezervasyon["TUR_BILGISI"], "tur_bitis": rezervasyon["REZ_BITIS"],
            "dogum_tarihi": rezervasyon.get("DOGUM_TARIHI"), "tc_pasaport": rezervasyon["TC_PASAPORT"],
            "kapasite": kapasite, "kayit": yeni_kayitli_sayisi,
            "musait": kalan_kontenjan, "iptal_token": rezervasyon["IPTAL_TOKEN"]
        }
        eposta_gonder(eposta_bilgileri)
    except Exception as e:
        current_app.logger.error(
            f"E-posta gönderimi sırasında bir hata oluştu: {e}")


def generate_sg_docx_from_template(rez: dict) -> str:
    """
    sahil_guvenlik.docx sablonunu açar ve:
    - 4. satira: Tur Adi + Tur Tarihi,
    - 8. madde tablosuna: 3. satirdan itibaren kayit yapanlar,
    - 9. program bölümüne: Tarih / Yeri / Saat(10:00)
    yazar. Geçici bir DOCX dosya yolu döndürür.
    """
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Sablon bulunamadi: {TEMPLATE_PATH}")

    form_data = _prepare_sg_form_data(rez)
    program_yeri = form_data["program_yeri"]
    program_tarih = form_data["program_tarih"]
    participants = form_data["participants"]

    # --- DOCX doldurma ---
    doc = Document(TEMPLATE_PATH)

    # 4) "kullanilacak parkur..." satirina: Tur Adi + Tur Tarihi
    parkur_text = program_yeri if not program_tarih else f"{program_yeri} {program_tarih}"
    for p in doc.paragraphs:
        if "4. kullanilacak parkur" in _norm_tr(p.text):
            _append_after_colon(p, parkur_text)
            break

    # 8) Dalgiç tablo: basliklari normalize ederek tabloyu bul
    tablo = _find_table_by_headers(
        doc, ["s.no", "adi", "soyadi", "belgesi", "pasaport", "milliyet", "yas"])
    if tablo:
        # 0 = baslik; 1 ve 2 = sabit satirlar; 3'ten itibaren dinamik
        start_row = 3
        needed = len(participants)
        while len(tablo.rows) < start_row + needed:
            tablo.add_row()
        for idx, d in enumerate(participants, start=start_row):
            r = tablo.rows[idx].cells
            r[0].text = str(idx)                 # S.NO (3,4,5,...)
            r[1].text = d["ad"]
            r[2].text = d["belgesi"]
            r[3].text = d["tc"]
            r[4].text = d["milliyet"]
            r[5].text = "" if d["yas"] == "" else str(d["yas"])

    # 9. bölüm
    for p in doc.paragraphs:
        t = _norm_tr(p.text)
        if t.startswith("dalis tarihi"):
            _append_after_colon(p, program_tarih)
        elif t.startswith("yeri"):
            _append_after_colon(p, program_yeri)
        elif t.startswith("saat"):
            _append_after_colon(p, "10:00")

    # Geçici docx kaydet
    tmpdir = tempfile.gettempdir()
    stamp = int(time.time() * 1000)
    docx_path = os.path.join(tmpdir, f"sg_preview_{stamp}.docx")
    doc.save(docx_path)
    return docx_path


def _normalize_ws(s: str) -> str:
    """Birden fazla boşluğu tek boşluğa indirir, baş/son boşlukları temizler."""
    return re.sub(r'\s+', ' ', (s or '').strip())


def parse_tur_bilgisi(tur_bilgisi: str):
    """
    Tamamen dinamik parser.
    Giriş örnekleri:
      - "KAŞ TEKNE TURU 05/09/2025 09:30"
      - "Bodrum Dalış 12.08.2025"
      - "Ayvalık 7-9-2025"
      - "Gökova Dalış Turu"  (tarih yok)
    Dönüş: (yer_ad, tarih_DD/MM/YYYY_ya_da_None, saat_ya_da_None)
    """
    text = _normalize_ws(tur_bilgisi)

    date_pat = re.compile(
        r'(?P<gun>\d{1,2})[./-](?P<ay>\d{1,2})[./-](?P<yil>\d{4})'
        r'(?:\s+(?P<saat>\d{1,2}:\d{2}))?',
        flags=re.IGNORECASE
    )

    match = None
    for m in date_pat.finditer(text):
        match = m  # son eşleşmeyi tercih et

    tarih_norm = None
    saat = None
    yer = text

    if match:
        gun = int(match.group('gun'))
        ay = int(match.group('ay'))
        yil = int(match.group('yil'))
        saat = match.group('saat')
        tarih_norm = f"{gun:02d}/{ay:02d}/{yil}"
        yer = (text[:match.start()] + text[match.end():]).strip()

    yer = _normalize_ws(yer)
    if not yer:
        yer = _normalize_ws(text)
    return yer, tarih_norm, saat


def _parse_date_multiformat(s: str) -> date | None:
    if not s:
        return None
    candidates = ['%d.%m.%Y', '%d/%m/%Y', '%d-%m-%Y', '%Y-%m-%d']
    txt = s.strip()
    for fmt in candidates:
        try:
            return datetime.strptime(txt, fmt).date()
        except Exception:
            continue
    return None


def calculate_age(dob_str: str, ref_date: date) -> int | str:
    dob = None
    if dob_str:
        norm = dob_str.replace('/', '.').replace('-', '.')
        dob = _parse_date_multiformat(norm)
    if not dob:
        return ""
    if not isinstance(ref_date, date):
        ref_date = date.today()
    return ref_date.year - dob.year - ((ref_date.month, ref_date.day) < (dob.month, dob.day))


@app.route('/sg_form_preview/<string:rez_id>', methods=['GET'])
def sg_form_preview(rez_id):
    if 'username' not in session:
        flash("Bu sayfayı görüntülemek için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    rez = customer_collection.find_one({"_id": ObjectId(rez_id)})
    if not rez:
        flash("Rezervasyon bulunamadı.", "danger")
        return redirect(url_for('tur_list'))

    form_data = _prepare_sg_form_data(rez)
    program_yeri = form_data["program_yeri"]
    program_tarih = form_data["program_tarih"]
    program_saat = form_data["program_saat"]
    dalgiclar = form_data["participants"]

    # Sıra numaraları 3’ten başlasın
    sira = 3
    for d in dalgiclar:
        d["sno"] = sira
        sira += 1

    context = {
        "belge_no": "1069",
        "yetki_belge": "1069",
        "malzeme_var_mi": "EVET",

        # 4. madde
        "parkur_tam_satir": f"4. kullanılacak parkur- tekne/yata ait biilgiler: {program_yeri}",
        "adi_satiri": "Adı:",
        "bayragi_satiri": "Bayrağı:",
        "baglama_limani_satiri": "Bağlama Limanı",

        # 5-6
        "sorumlular_baslik": "5. Dalış Yapacak Grubun Sorumluları:",
        "sorumlu_turk_baslik": "6. Sorumlu Türk Dalış Rehberinin:",
        "sorumlular_satiri": "Adı, Soyadı: BİLGE ŞEN, ERCÜMENT ŞEN",
        "ikinci_adi_satiri": "Adı, Soyadı: ERCÜMENT ŞEN",
        "sorumlu_milliyet": "TÜRKİYE CUMHURİYETİ",
        "belge_yeri": "TSSF",
        "dalicilik_yeri": "TSSF",

        # 7
        "yedi_baslik": "7. Dalış Şekli:",
        "turkiye_acente": "Türkiye'de Bağlı olduğu Acente:",
        "techizatsiz": "Techizatsız",
        "turk_acente_sorumlu": "Türk Acente Sorumlusunun Adı/Soyadı:",
        "techizatli_x": "Techizatlı   X",

        # 8 – sadece dinamik satırlar gönderiyoruz
        "dalgiclar_dinamik": dalgiclar,

        # 9
        "dokuz_baslik": "9.Dalış Programı:",
        "program_tarih": program_tarih or "",
        "program_yeri": program_yeri,
        "program_saat": program_saat,

        "internal": "HİZMETE ÖZEL - INTERNAL",
    }

    return render_template("sg_form_preview.html", **context)


# KVKK PDF dosyasını ilk seferde yüklemek için (manuel bir kere çalıştırın)
def pdf_yukle_ve_kaydet():
    pdf_path = "static/ACIK RIZA METNI_VERTICAL.pdf"
    if os.path.exists(pdf_path):
        with open(pdf_path, "rb") as f:
            pdf_base64 = base64.b64encode(f.read()).decode("utf-8")
        if not kvkk_pdf_koleksiyon.find_one({"VERSIYON": "v1.2"}):
            kvkk_pdf_koleksiyon.insert_one({
                "DOSYA_ADI": "ACIK RIZA METNI_VERTICAL.pdf",
                "VERSIYON": "v1.2",
                "TUR": "application/pdf",
                "BASE64": pdf_base64
            })
        print("PDF başarıyla base64 olarak MongoDB'ye yüklendi.")
    else:
        print("PDF dosyası static klasöründe bulunamadı!")


# Update the `hesapla_pdf_sha256` function to accept a file path as an argument and calculate the hash.
def hesapla_pdf_sha256(dosya_yolu):
    with open(dosya_yolu, 'rb') as f:
        pdf_bytes = f.read()
        return hashlib.sha256(pdf_bytes).hexdigest()


# MongoDB'deki base64 PDF'den hash hesaplama
kvkk_doc = db["kvkk_pdf"].find_one({"VERSIYON": "v1.2"})
if kvkk_doc and kvkk_doc.get("BASE64"):
    base64_str = kvkk_doc["BASE64"]
    pdf_bytes = base64.b64decode(base64_str)
    hash_degeri = hashlib.sha256(pdf_bytes).hexdigest()
    print("PDF Hash:", hash_degeri)
else:
    hash_degeri = None  # veya raise Exception("KVKK PDF bulunamadı")


@app.route("/kvkk/view")
def kvkk_pdf_goster():
    doc = kvkk_pdf_koleksiyon.find_one({"VERSIYON": "v1.2"})
    if not doc:
        return "KVKK metni bulunamadı.", 404
    base64_data = doc["BASE64"]
    return f'''
    <!DOCTYPE html>
    <html lang="tr">
    <head>
        <meta charset="UTF-8">
        <title>KVKK Metni</title>
    </head>
    <body>
        <h3>KVKK Açık Rıza Metni</h3>
        <iframe src="data:application/pdf;base64,{base64_data}" width="100%" height="700px"></iframe>
    </body>
    </html>
    '''


@app.route('/static/<path:filename>')
def static_files(filename):
    return send_from_directory(app.static_folder, filename)


@app.route('/sg_form_pdf/<string:rez_id>', methods=['GET'])
def sg_form_pdf(rez_id):
    if 'username' not in session:
        flash("Bu sayfayı görüntülemek için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    rez = customer_collection.find_one({"_id": ObjectId(rez_id)})
    if not rez:
        flash("Rezervasyon bulunamadı.", "danger")
        return redirect(url_for('tur_list'))

    # sg_form_preview ile aynı mantığı kullanarak context oluştur
    form_data = _prepare_sg_form_data(rez)
    dalgiclar = form_data["participants"]
    sira = 3
    for d in dalgiclar:
        d["sno"] = sira
        sira += 1

    context = {
        "belge_no": "1069",
        "yetki_belge": "1069",
        "malzeme_var_mi": "EVET",
        "parkur_tam_satir": f"4. kullanılacak parkur- tekne/yata ait biilgiler: {form_data['program_yeri']}",
        "adi_satiri": "Adı:", "bayragi_satiri": "Bayrağı:", "baglama_limani_satiri": "Bağlama Limanı",
        "sorumlular_baslik": "5. Dalış Yapacak Grubun Sorumluları:",
        "sorumlu_turk_baslik": "6. Sorumlu Türk Dalış Rehberinin:",
        "sorumlular_satiri": "Adı, Soyadı: BİLGE ŞEN, ERCÜMENT ŞEN",
        "ikinci_adi_satiri": "Adı, Soyadı: ERCÜMENT ŞEN",
        "sorumlu_milliyet": "TÜRKİYE CUMHURİYETİ", "belge_yeri": "TSSF", "dalicilik_yeri": "TSSF",
        "yedi_baslik": "7. Dalış Şekli:", "turkiye_acente": "Türkiye'de Bağlı olduğu Acente:",
        "techizatsiz": "Techizatsız", "turk_acente_sorumlu": "Türk Acente Sorumlusunun Adı/Soyadı:",
        "techizatli_x": "Techizatlı   X",
        "dalgiclar_dinamik": dalgiclar,
        "dokuz_baslik": "9.Dalış Programı:",
        "program_tarih": form_data["program_tarih"],
        "program_yeri": form_data["program_yeri"],
        "program_saat": form_data["program_saat"],
        "internal": "HİZMETE ÖZEL - INTERNAL",
    }

    try:
        # HTML içeriğini 'sg_form_preview.html' şablonundan oluştur
        html_content = render_template('sg_form_preview.html', **context)

        # --- PDF'ten butonu gizlemek için özel CSS ---
        # .topbar içindeki butonu (<a> veya <button>) gizle, başlığı koru
        css_content = ".topbar a, .topbar button { display: none !important; }"
        css_path = None
        pdf_bytes = None

        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.css', encoding='utf-8') as css_file:
                css_file.write(css_content)
                css_path = css_file.name

            options = {
                'page-size': 'A4',
                'margin-top': '0.75in', 'margin-right': '0.75in',
                'margin-bottom': '0.75in', 'margin-left': '0.75in',
                'encoding': "UTF-8",
                'enable-local-file-access': None,
                'user-style-sheet': css_path  # Kendi CSS dosyamızı ekliyoruz
            }
            if not config:
                raise RuntimeError(
                    "wkhtmltopdf yapılandırılamadı. Lütfen WKHTMLTOPDF_PATH tanımlayın.")
            if not config:
                raise RuntimeError(
                    "wkhtmltopdf yapılandırılamadı. Lütfen WKHTMLTOPDF_PATH tanımlayın.")
            pdf_bytes = pdfkit.from_string(
                html_content, False, configuration=config, options=options)
        finally:
            if css_path and os.path.exists(css_path):
                os.remove(css_path)
        # --- Bitiş: özel CSS ---

        return send_file(
            BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=False,  # Tarayıcıda göstermek için False
            download_name="sahil_guvenlik_onizleme.pdf"
        )
    except Exception as e:
        current_app.logger.error(f"Tekil SG Formu PDF oluşturma hatası: {e}")
        flash(f"PDF oluşturulurken bir hata oluştu: {e}", "danger")
        return redirect(url_for('tur_list'))


@app.route('/sg_form_word/<string:rez_id>')
def sg_form_word(rez_id):
    if 'username' not in session:
        flash("Bu işlemi yapmak için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    rez = customer_collection.find_one({"_id": ObjectId(rez_id)})
    if not rez:
        flash("Rezervasyon bulunamadı.", "danger")
        return redirect(url_for('tur_list'))

    docx_path = ""
    try:
        # 1) DOCX oluştur
        docx_path = generate_sg_docx_from_template(rez)

        # 2) Belleğe alıp gönder
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

        # Dosya adını oluştur
        adsoyad = rez.get("AD", "bilinmiyor").replace(' ', '_')
        tur_bilgisi = rez.get("TUR_BILGISI", "")
        _, tarih_norm, _ = parse_tur_bilgisi(tur_bilgisi)
        tarih_str = tarih_norm or rez.get("TARIH", "tarih_yok")

        dosya_adi = f"sahil_guvenlik_{adsoyad}_{tarih_str.replace('/', '-')}.docx"

        return send_file(
            BytesIO(docx_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=dosya_adi
        )
    except Exception as e:
        current_app.logger.error(f"SG Formu DOCX oluşturulurken hata: {e}")
        flash("Form oluşturulurken bir hata oluştu.", "danger")
        return redirect(url_for('tur_list'))
    finally:
        # Geçici dosyayı sil
        if docx_path and os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except Exception as e:
                current_app.logger.error(
                    f"Geçici DOCX dosyası silinemedi: {e}")


@app.route('/sg_form_tur_preview/<string:tur_id>')
def sg_form_tur_preview(tur_id):
    """
    Bir tura kayıtlı tüm katılımcılar için birleştirilmiş bir Sahil Güvenlik formu önizlemesi (HTML) oluşturur.
    """
    if 'username' not in session:
        flash("Bu işlemi yapmak için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    try:
        tur_object_id = ObjectId(tur_id)
    except Exception:
        flash("Geçersiz tur ID'si.", "danger")
        return redirect(url_for('tur_list'))

    # 1. Tur bilgilerini al
    tur = tur_collection.find_one({"_id": tur_object_id})
    if not tur:
        flash("Tur bulunamadı.", "danger")
        return redirect(url_for('tur_list'))

    # 2. Bu tura ait tüm rezervasyonları bul
    rezervasyonlar = list(customer_collection.find({"TUR_ID": tur_object_id}))
    if not rezervasyonlar:
        flash("Bu tur için henüz rezervasyon yapılmamış.", "warning")
        return redirect(url_for('tur_list'))

    # 3. Tüm katılımcıları tek bir "dalgıçlar" listesinde topla
    all_dalgiclar_raw = []
    for rez in rezervasyonlar:
        if isinstance(rez.get("DALGICLAR"), list) and rez["DALGICLAR"]:
            all_dalgiclar_raw.extend(rez["DALGICLAR"])
        else:
            all_dalgiclar_raw.append({
                "AD": rez.get("AD"),
                "SERTIFIKA": rez.get("SERTIFIKA"),
                "TC_PASAPORT": rez.get("TC_PASAPORT"),
                "MILLIYET": rez.get("MILLIYET"),
                "DOGUM_TARIHI": rez.get("DOGUM_TARIHI"),
            })

    # 4. Form oluşturucu fonksiyona göndermek için ana bir "rezervasyon" nesnesi oluştur
    master_rez = {
        "TUR_BILGISI": tur.get("TUR_ISMI", "") + " " + tur.get("BASLANGIC_TARIHI").strftime('%d/%m/%Y'),
        "TARIH": tur.get("BASLANGIC_TARIHI").strftime('%d/%m/%Y'),
        "TARIH_NESNESI": tur.get("BASLANGIC_TARIHI"),
        "DALGICLAR": all_dalgiclar_raw,
        "MILLIYET": "T.C."  # Varsayılan milliyet
    }

    # 5. Verileri _prepare_sg_form_data ile işle
    form_data = _prepare_sg_form_data(master_rez)

    # 6. Template'e göndermek için verileri hazırla
    context = {
        "tur_id": tur_id,
        "parkur_ve_tarih": f'{form_data["program_yeri"]} {form_data["program_tarih"]}',
        "dalis_tarihi": form_data["program_tarih"],
        "dalis_yeri": form_data["program_yeri"],
        "dalis_saat": form_data["program_saat"],
        "participants": form_data["participants"]
    }

    return render_template('sg.html', **context)


@app.route('/sg_form_tur_pdf/<string:tur_id>')
def sg_form_tur_pdf(tur_id):
    """
    sg.html önizleme sayfasını bellekte oluşturur, bunu bir PDF dosyasına dönüştürür
    ve indirme için kullanıcıya gönderir.
    """
    if 'username' not in session:
        flash("Bu işlemi yapmak için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    try:
        tur_object_id = ObjectId(tur_id)
    except Exception:
        flash("Geçersiz tur ID'si.", "danger")
        return redirect(url_for('tur_list'))

    # 1. Tur bilgilerini ve katılımcıları al (sg_form_tur_preview ile aynı mantık)
    tur = tur_collection.find_one({"_id": tur_object_id})
    if not tur:
        flash("Tur bulunamadı.", "danger")
        return redirect(url_for('tur_list'))

    rezervasyonlar = list(customer_collection.find({"TUR_ID": tur_object_id}))
    if not rezervasyonlar:
        flash("Bu tur için henüz rezervasyon yapılmamış.", "warning")
        return redirect(url_for('tur_list'))

    all_dalgiclar_raw = []
    for rez in rezervasyonlar:
        if isinstance(rez.get("DALGICLAR"), list) and rez["DALGICLAR"]:
            all_dalgiclar_raw.extend(rez["DALGICLAR"])
        else:
            all_dalgiclar_raw.append({
                "AD": rez.get("AD"),
                "SERTIFIKA": rez.get("SERTIFIKA"),
                "TC_PASAPORT": rez.get("TC_PASAPORT"),
                "MILLIYET": rez.get("MILLIYET"),
                "DOGUM_TARIHI": rez.get("DOGUM_TARIHI"),
            })

    master_rez = {
        "TUR_BILGISI": tur.get("TUR_ISMI", "") + " " + tur.get("BASLANGIC_TARIHI").strftime('%d/%m/%Y'),
        "TARIH": tur.get("BASLANGIC_TARIHI").strftime('%d/%m/%Y'),
        "TARIH_NESNESI": tur.get("BASLANGIC_TARIHI"),
        "DALGICLAR": all_dalgiclar_raw,
        "MILLIYET": "T.C."
    }

    form_data = _prepare_sg_form_data(master_rez)

    context = {
        "tur_id": tur_id,
        "parkur_ve_tarih": f'{form_data["program_yeri"]} {form_data["program_tarih"]}',
        "dalis_tarihi": form_data["program_tarih"],
        "dalis_yeri": form_data["program_yeri"],
        "dalis_saat": form_data["program_saat"],
        "participants": form_data["participants"]
    }

    try:
        # 2. HTML içeriğini bir string olarak oluştur
        html_content = render_template('sg.html', **context)

        # --- PDF'ten butonu gizlemek için özel CSS ---
        # .topbar içindeki butonu (<a> veya <button>) gizle, başlığı koru
        css_content = ".topbar a, .topbar button { display: none !important; }"
        css_path = None
        pdf_bytes = None

        try:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.css', encoding='utf-8') as css_file:
                css_file.write(css_content)
                css_path = css_file.name

            options = {
                'page-size': 'A4',
                'margin-top': '0.75in', 'margin-right': '0.75in',
                'margin-bottom': '0.75in', 'margin-left': '0.75in',
                'encoding': "UTF-8",
                'enable-local-file-access': None,
                'user-style-sheet': css_path  # Kendi CSS dosyamızı ekliyoruz
            }
            pdf_bytes = pdfkit.from_string(
                html_content, False, configuration=config, options=options)
        finally:
            if css_path and os.path.exists(css_path):
                os.remove(css_path)
        # --- Bitiş: özel CSS ---

        # 4. Dosya adını oluştur ve kullanıcıya gönder
        tur_ismi = tur.get("TUR_ISMI", "tur").replace(' ', '_')
        tarih_str = tur.get("BASLANGIC_TARIHI").strftime('%d-%m-%Y')
        dosya_adi = f"sahil_guvenlik_formu_{tur_ismi}_{tarih_str}.pdf"

        return send_file(
            BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=dosya_adi
        )
    except Exception as e:
        current_app.logger.error(f"HTML'den PDF oluşturma hatası: {e}")
        flash(f"Form indirilirken bir hata oluştu: {e}", "danger")
        return redirect(url_for('tur_list'))


@app.route('/yeni_tur', methods=['GET', 'POST'])
def yeni_tur():
    # 1. Yetkilendirme Kontrolü
    if 'username' not in session:
        flash("Bu sayfayı görüntülemek için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    if request.method == 'POST':
        tur_ismi = request.form.get('tur_ismi')
        baslangic = request.form.get('baslangic_tarihi')
        bitis = request.form.get('bitis_tarihi')
        detaylar = request.form.get('detaylar')
        fiyat = request.form.get('fiyat')  # New field for price

        # 2. Veri Doğrulama ve Hata Yönetimi
        try:
            kontenjan = int(request.form.get('kontenjan'))
            baslangic_dt = datetime.strptime(baslangic, "%d/%m/%Y")
            bitis_dt = datetime.strptime(bitis, "%d/%m/%Y")
            fiyat = float(fiyat)  # Ensure the price is a valid float

            # 3. Mantıksal Kontroller
            if bitis_dt < baslangic_dt:
                flash("Bitiş tarihi, başlangıç tarihinden önce olamaz.", "danger")
                return render_template("yeni_tur.html", form_data=request.form)
            if kontenjan <= 0:
                flash("Kontenjan pozitif bir sayı olmalıdır.", "danger")
                return render_template("yeni_tur.html", form_data=request.form)
            if fiyat <= 0:
                flash("Fiyat pozitif bir sayı olmalıdır.", "danger")
                return render_template("yeni_tur.html", form_data=request.form)
        except (ValueError, TypeError):
            flash(
                "Lütfen kontenjan ve fiyat için geçerli bir sayı ve tarihler için doğru format girin.", "danger")
            return render_template("yeni_tur.html", form_data=request.form)

        tur = {
            "TUR_ISMI": tur_ismi,
            "BASLANGIC_TARIHI": baslangic_dt,
            "BITIS_TARIHI": bitis_dt,
            "KONTENJAN": kontenjan,
            "DETAYLAR": detaylar,
            "FIYAT": fiyat,
            "KAYIT_TARIHI": datetime.now()
        }

        tur_collection.insert_one(tur)
        flash("Yeni tur başarıyla eklendi!", "success")
        return redirect(url_for('tur_list'))

    return render_template("yeni_tur.html", form_data={})


@app.route('/admin', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        # MongoDB'de kullanıcıyı kullanıcı adına göre bul
        user = users_collection.find_one({"username": username})

        # Kullanıcı varsa ve girilen şifre, veritabanındaki hash ile eşleşiyorsa
        # Kullnici hesabi olusturmadan once hash_generator.py dosyasını çalıştırarak hash oluşturulmalı
        # ve DB icerisindeki users koleksiyonuna eklenmelidir.
        if user and check_password_hash(user['password'], password):
            session['username'] = username
            print("Kullanıcı oturumu açıldı:", session)
            return redirect(url_for('rezervasyon_list'))
        else:
            flash("Hatalı kullanıcı adı veya şifre", "danger")
            return redirect(url_for('login'))

    return render_template('rez_admin_vertical.html')


@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))


@app.route('/cikti_al', methods=['GET'])
def cikti_al():
    # 1. Yetkilendirme Kontrolü
    if 'username' not in session:
        flash("Bu işlemi yapmak için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    # 2. Sadece bugünün ve ileri tarihli rezervasyonları al
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)

    # 3. Verileri yardımcı fonksiyon ile çek
    match_filter = {'TARIH_NESNESI': {'$gte': today_start}}
    rezervasyon_listesi = _get_rezervasyonlar(match_filter)

    # Verileri bir pandas DataFrame'e dönüştür
    data = []
    for rezervasyon in rezervasyon_listesi:
        tur_bilgisi = rezervasyon.get("TUR_BILGISI", "Günlük Dalış")
        data.append({
            "Ad": rezervasyon.get("AD", ""),
            "Telefon": rezervasyon.get("GSM", ""),
            "E-posta": rezervasyon.get("EMAIL", ""),
            "TC/Pasaport No": rezervasyon.get("TC_PASAPORT", ""),
            "Tur/Tarih": f"{tur_bilgisi} ({rezervasyon.get('TARIH', '')})",
            "Sertifika": rezervasyon.get("SERTIFIKA", ""),
            "Kayıt Tarihi": rezervasyon.get("KAYIT_TARIHI", ""),
            "Kredi": rezervasyon.get("kredi_bilgisi_arr", {}).get("KREDI", "Kredi Yok")
        })

    df = pd.DataFrame(data)

    # "Sıra" sütununu ekle
    df.index += 1  # Sıra numaraları 1'den başlasın
    df.reset_index(inplace=True)
    df.rename(columns={"index": "Sıra"}, inplace=True)

    # Excel dosyasını oluştur
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Rezervasyonlar")

    # Dosyayı kullanıcıya indirme olarak sun
    output.seek(0)
    response = make_response(output.read())
    response.headers["Content-Disposition"] = "attachment; filename=rezervasyonlar.xlsx"
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return response


def temizle_rezervasyonlar():
    # Istanbul saatine göre bugünün başlangıcını al
    istanbul_tz = pytz.timezone("Europe/Istanbul")
    # Zaman dilimi bilgisiyle birlikte bugünün başlangıcını al
    today_start_aware = datetime.now(istanbul_tz).replace(
        hour=0, minute=0, second=0, microsecond=0)

    # Veritabanındaki tarih (TARIH_NESNESI) naive (zaman dilimi bilgisi olmadan)
    # saklandığı için, karşılaştırma yapmadan önce zaman dilimi bilgisini kaldırıyoruz.
    today_start_naive = today_start_aware.replace(tzinfo=None)

    # Rezervasyon Tarihi bugünden eski olan kayıtları bul (naive datetime ile karşılaştırma)
    arsivlenecek_kayitlar = list(customer_collection.find(
        {"TARIH_NESNESI": {"$lt": today_start_naive}}))

    if arsivlenecek_kayitlar:
        for kayit in arsivlenecek_kayitlar:
            kayit["ARSIV_TARIHI"] = datetime.now()
        try:
            db["gecmis_rezervasyon"].insert_many(arsivlenecek_kayitlar)
            customer_collection.delete_many(
                {"_id": {"$in": [k["_id"] for k in arsivlenecek_kayitlar]}})
            print(f"{len(arsivlenecek_kayitlar)} kayıt arşiv koleksiyonuna taşındı.")
        except Exception as e:
            print(f"Arşivleme sırasında hata oluştu: {e}")
    else:
        print("Arşivlenecek geçmiş rezervasyon bulunamadı.")


# APScheduler ile zamanlanmış görev
scheduler = BackgroundScheduler(timezone="Europe/Istanbul")
scheduler.add_job(
    func=temizle_rezervasyonlar,
    trigger="cron",
    hour=23,
    minute=0,
    id="cleanup_old_reservations",
    replace_existing=True
)
should_start = True
if os.environ.get("WERKZEUG_RUN_MAIN") == "true" or os.environ.get("RUN_MAIN") == "true":
    should_start = True  # gerçek process
elif os.environ.get("WERKZEUG_RUN_MAIN") is None:
    # Üretimde / tek process'te
    should_start = True

if should_start and not scheduler.running:
    scheduler.start()


@app.route('/rezervasyon_list')
def rezervasyon_list():
    if 'username' not in session:
        flash("Lütfen giriş yapın.", "danger")

    current_year = datetime.now().year
    min_year = _min_year_from_db()
    # [2025, 2024, 2023, …]
    years = list(range(current_year, min_year - 1, -1))

    # Filtreleme için temel $match sorgusunu oluştur
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    match_filter = {'TARIH_NESNESI': {'$gte': today_start}}

    # Verileri veritabanından çekmek için projeksiyon (projection) tanımla
    projection = {
        '_id': {'$toString': '$_id'},
        'AD': '$AD',
        'GSM': '$GSM',
        'EMAIL': {'$ifNull': ['$EMAIL', 'E-posta Yok']},
        'TC_PASAPORT': '$TC_PASAPORT',
        'DOGUM_TARIHI': '$DOGUM_TARIHI',
        'SERTIFIKA': '$SERTIFIKA',
        'TARIH': '$TARIH',
        # Template'de `TUR_ISMI` kullanıldığı için alanı ona göre yeniden adlandırıyoruz.
        'TUR_ISMI': {'$ifNull': ['$TUR_BILGISI', 'Günlük Dalış']},
        'KAYIT_TARIHI': '$KAYIT_TARIHI',
        'KREDI': {'$ifNull': ['$kredi_bilgisi_arr.KREDI', 'Kredi Yok']}
    }

    # Yardımcı fonksiyonu kullanarak verileri çek
    rezervasyon_listesi = _get_rezervasyonlar(match_filter, projection)

    return render_template('rezervasyon_list_vertical.html', current_year=current_year, years=years, rezervasyonlar=rezervasyon_listesi)


@app.route('/turlar')
def tur_list():
    if 'username' not in session:
        flash("Lütfen giriş yapın.", "danger")
        return redirect(url_for('login'))

    current_year = datetime.now().year
    min_year = _min_year_from_db()
    # [2025, 2024, 2023, …]
    years = list(range(current_year, min_year - 1, -1))

    # Sadece bugünün ve gelecekteki turları göstermek için günün başlangıcını al
    today_start = datetime.now().replace(
        hour=0, minute=0, second=0, microsecond=0)

    # Aggregation pipeline ile turları, rezervasyonları ve kayıt sayılarını birleştir
    pipeline = [
        {
            '$match': {
                "BITIS_TARIHI": {"$gte": today_start}
            }
        },
        {
            '$lookup': {
                'from': 'Rezervasyon',
                'localField': '_id',
                'foreignField': 'TUR_ID',
                'as': 'rezervasyonlar'
            }
        },
        {
            '$addFields': {
                'kayit_sayisi': {'$size': '$rezervasyonlar'}
            }
        },
        {
            '$sort': {'BASLANGIC_TARIHI': 1}
        }
    ]
    turlar_ile_rezervasyonlar = list(tur_collection.aggregate(pipeline))

    return render_template('tur_list_vertical.html', current_year=current_year, years=years, turlar=turlar_ile_rezervasyonlar)
# REZERVASYON DUZENLE


@app.route('/tur_duzenle/<string:tur_id>', methods=['GET', 'POST'])
def tur_duzenle(tur_id):
    if 'username' not in session:
        flash("Lütfen giriş yapın.", "danger")
        return redirect(url_for('login'))

    tur = tur_collection.find_one({"_id": ObjectId(tur_id)})
    if not tur:
        flash("Düzenlenecek tur bulunamadı.", "danger")
        return redirect(url_for('tur_list'))

    if request.method == 'POST':
        # Formdan gelen verileri al
        tur_ismi = request.form.get('tur_ismi')
        baslangic = request.form.get('baslangic_tarihi')
        bitis = request.form.get('bitis_tarihi')
        detaylar = request.form.get('detaylar')
        fiyat = request.form.get('fiyat')  # Ensure the price is fetched

        # Veri Doğrulama ve Hata Yönetimi
        try:
            kontenjan = int(request.form.get('kontenjan'))
            baslangic_dt = datetime.strptime(baslangic, "%d/%m/%Y")
            bitis_dt = datetime.strptime(bitis, "%d/%m/%Y")
            fiyat = float(fiyat)  # Ensure the price is a valid float

            # Mantıksal Kontroller
            if bitis_dt < baslangic_dt:
                flash("Bitiş tarihi, başlangıç tarihinden önce olamaz.", "danger")
                return render_template('tur_duzenle.html', tur=tur)
            if kontenjan <= 0:
                flash("Kontenjan pozitif bir sayı olmalıdır.", "danger")
                return render_template('tur_duzenle.html', tur=tur)
            if fiyat <= 0:
                flash("Fiyat pozitif bir sayı olmalıdır.", "danger")
                return render_template('tur_duzenle.html', tur=tur)
        except (ValueError, TypeError):
            flash(
                "Lütfen kontenjan ve fiyat için geçerli bir sayı ve tarihler için doğru format girin.", "danger")
            return render_template('tur_duzenle.html', tur=tur)

        # Güncelleme verisini oluştur
        guncelleme = {
            "$set": {
                "TUR_ISMI": tur_ismi, "BASLANGIC_TARIHI": baslangic_dt,
                "BITIS_TARIHI": bitis_dt, "KONTENJAN": kontenjan, "DETAYLAR": detaylar, "FIYAT": fiyat
            }
        }
        tur_collection.update_one({"_id": ObjectId(tur_id)}, guncelleme)
        flash("Tur başarıyla güncellendi.", "success")
        return redirect(url_for('tur_list'))

    # GET isteği için, formu tur verileriyle doldurarak göster
    return render_template('tur_duzenle.html', tur=tur)


@app.route('/rezervasyon_duzenle/<string:id>', methods=['GET', 'POST'])
def rezervasyon_duzenle(id):
    if 'username' not in session:
        flash("Bu işlemi yapmak için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))
    # İlgili rezervasyonu MongoDB'den çek
    rezervasyon = customer_collection.find_one({"_id": ObjectId(id)})
    print(rezervasyon)

    if not rezervasyon:
        flash("Düzenlenecek rezervasyon bulunamadı.", "danger")
        return redirect(url_for('rezervasyon_list'))

    if request.method == 'POST':
        yeni_tarih_str = request.form.get('tarih')
        try:
            # 2. Tarih nesnesini de güncelle
            yeni_tarih_nesnesi = datetime.strptime(yeni_tarih_str, '%d/%m/%Y')
        except ValueError:
            flash("Geçersiz tarih formatı. Lütfen GG/AA/YYYY formatında girin.", "danger")
            return redirect(url_for('rezervasyon_duzenle', id=id))

        # Formdan gelen verileri al ve MongoDB'de güncelle
        updated_data = {
            "AD": request.form.get('ad'),
            "GSM": request.form.get('gsm'),
            "EMAIL": request.form.get('email'),
            "TC_PASAPORT": request.form.get('tc_pasaport'),
            "DOGUM_TARIHI": request.form.get('dogum_tarihi'),
            "SERTIFIKA": request.form.get('sertifika'),
            "TARIH": yeni_tarih_str,
            "TARIH_NESNESI": yeni_tarih_nesnesi,
            "KAYIT_TARIHI": rezervasyon["KAYIT_TARIHI"]
        }
        customer_collection.update_one(
            {"_id": ObjectId(id)}, {"$set": updated_data})
        flash('Rezervasyon başarıyla güncellendi.', 'success')
        return redirect(url_for('rezervasyon_list'))

    # Düzenleme sayfasına render
    return render_template('rez_duzenle_vertical.html', rezervasyon=rezervasyon)


# REZERVASYON SIL
@app.route('/rezervasyon_sil/<string:id>', methods=['POST'])
def rezervasyon_sil(id):
    # 1. Yetkilendirme Kontrolü
    if 'username' not in session:
        flash("Bu işlemi yapmak için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    try:
        # 2. Silme işlemini gerçekleştir ve sonucu al
        result = customer_collection.delete_one({"_id": ObjectId(id)})

        # 3. Sonucu kontrol et
        if result.deleted_count > 0:
            flash('Rezervasyon başarıyla silindi.', 'success')
        else:
            flash('Silinecek rezervasyon bulunamadı.', 'warning')
    except Exception as e:
        flash(f'Rezervasyon silinirken bir hata oluştu: {e}', 'danger')

    return redirect(url_for('rezervasyon_list'))


@app.route('/rezervasyon_iptal/<string:token>', methods=['GET', 'POST'])
def rezervasyon_iptal_et(token):
    # Token'a göre rezervasyonu bul
    rezervasyon = customer_collection.find_one({"IPTAL_TOKEN": token})

    # Eğer rezervasyon bulunamazsa (geçersiz token veya zaten iptal edilmiş)
    if not rezervasyon:
        mesaj = "Geçersiz veya daha önce kullanılmış bir iptal bağlantısı."
        return render_template("rezervasyon_onay_vertical.html", mesaj=mesaj)

    # Kullanıcı "Evet, İptal Et" butonuna tıkladığında
    if request.method == 'POST':
        customer_collection.delete_one({"IPTAL_TOKEN": token})
        mesaj = "Rezervasyonunuz başarıyla iptal edilmiştir."
        # İsteğe bağlı: Yöneticiye ve kullanıcıya iptal edildiğine dair e-posta gönderilebilir.
        return render_template("rezervasyon_onay_vertical.html", mesaj=mesaj)

    # GET isteği için onay sayfasını göster
    return render_template("rezervasyon_onay_vertical.html", iptal_token=token)


# E-posta gönderme fonksiyonu
# email.env dosyasındaki ortam değişkenlerini yükle
load_dotenv(os.path.join(os.path.dirname(__file__), 'email.env'))


def eposta_gonder(bilgiler):
    # 1. Hassas bilgileri ortam değişkenlerinden al
    sender_email = os.getenv("SENDER_EMAIL")
    # Gmail uygulama şifresi kullanılmalı
    password = os.getenv("EMAIL_APP_PASSWORD")
    receiver_email = os.getenv("ADMIN_EMAIL")

    if not password:
        current_app.logger.error(
            "HATA: E-posta şifresi (EMAIL_APP_PASSWORD) ortam değişkenlerinde tanımlanmamış.")
        return

    try:
        # 2. SMTP sunucusuna bağlan
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(sender_email, password)

            # --- Yöneticiye mail gönder ---
            try:
                # 3. HTML template kullanarak daha zengin mail içeriği oluştur
                admin_html = render_template(
                    'email/admin_notification.html', bilgiler=bilgiler)
                message_admin = MIMEMultipart("alternative")
                message_admin["Subject"] = f"Yeni Rezervasyon: {bilgiler.get('adsoyad', 'Bilinmiyor')}"
                message_admin["From"] = sender_email
                message_admin["To"] = receiver_email
                message_admin.attach(MIMEText(admin_html, "html"))
                server.sendmail(sender_email, receiver_email,
                                message_admin.as_string())
                current_app.logger.info(
                    "Yönetici bilgilendirme maili gönderildi.")
            except Exception as e:
                current_app.logger.error(
                    f"Yöneticiye mail gönderilirken hata oluştu: {e}")

            # --- Kullanıcıya mail gönder ---
            user_email = bilgiler.get("email")
            if user_email:
                try:
                    user_html = render_template(
                        'email/user_confirmation.html', bilgiler=bilgiler)
                    message_user = MIMEMultipart("alternative")
                    message_user["Subject"] = "Rezervasyon Onayınız – Vertical Dalış Merkezi"
                    message_user["From"] = sender_email
                    message_user["To"] = user_email
                    message_user.attach(MIMEText(user_html, "html"))
                    server.sendmail(sender_email, user_email,
                                    message_user.as_string())
                    current_app.logger.info(
                        f"Kullanıcıya ({user_email}) onay maili gönderildi.")
                except Exception as e:
                    current_app.logger.error(
                        f"Kullanıcıya mail gönderilirken hata oluştu: {e}")

    except smtplib.SMTPAuthenticationError:
        current_app.logger.error(
            "SMTP kimlik doğrulama hatası. Lütfen e-posta ve şifreyi kontrol edin.")
    except Exception as e:
        current_app.logger.error(
            f"E-posta gönderiminde genel bir hata oluştu: {e}")


@app.route('/')
def index():
    # 1. Karşılaştırma için günün başlangıcını al
    today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    # 2. Sadece aktif turları al ve başlangıç tarihine göre sırala
    turlar = list(tur_collection.find(
        {"BITIS_TARIHI": {"$gte": today_start}}
    ).sort("BASLANGIC_TARIHI", 1))
    return render_template("rez_index_vertical.html", turlar=turlar)


# Update the `/submit` route to correctly use `url_for` for the `/odeme` endpoint.
@app.route('/submit', methods=['POST'])
def submit():
    # Form verilerini al
    tc_pasaport = request.form.get('tc_pasaport', '').strip()
    adsoyad = request.form.get('adsoyad', '').strip()
    telefon = request.form.get('telefon', '').strip()
    email = request.form.get('email', '').strip()
    dogum_tarihi = request.form.get('dogumtarihi')
    adres = request.form.get('adres')
    sertifika = request.form.get('sertifika')
    tur_id = request.form.get('tur_hidden')

    kvkk_onay = 'kvkk_onay' in request.form
    kvkk_onay_tarihi = (request.form.get('kvkk_onay_tarihi') or '').strip()
    kvkk_doc = kvkk_pdf_koleksiyon.find_one(sort=[("VERSIYON", -1)])
    kvkk_versiyon = kvkk_doc["VERSIYON"] if kvkk_doc else "Bilinmiyor"
    kvkk_ip_adresi = request.remote_addr

    if not all([tc_pasaport, adsoyad, telefon, email, tur_id]):
        flash("Lütfen tüm zorunlu alanları doldurun ve bir tur seçin.", "danger")
        return redirect('/')

    if not kvkk_onay:
        flash("Rezervasyona devam etmek için KVKK metnini onaylamalısınız.", "warning")
        return redirect('/')

    # Tur bilgisi
    secili_tur = tur_collection.find_one({"_id": ObjectId(tur_id)})
    if not secili_tur:
        flash("Geçersiz bir tur seçildi. Lütfen tekrar deneyin.", "danger")
        return redirect('/')

    # Kontenjan kontrolü
    kayitli_sayisi = customer_collection.count_documents(
        {"TUR_ID": ObjectId(tur_id)})
    if kayitli_sayisi >= secili_tur.get("KONTENJAN", 0):
        flash(
            f"'{secili_tur.get('TUR_ISMI')}' turu için kontenjan dolmuştur.", "danger")
        return redirect('/')

    # Aynı tur için mükerrer kayıt kontrolü
    mevcut_rezervasyon = customer_collection.find_one(
        {"TC_PASAPORT": tc_pasaport, "TUR_ID": ObjectId(tur_id)})
    if mevcut_rezervasyon:
        flash("Bu tur için zaten bir rezervasyonunuz bulunmaktadır.", "danger")
        return redirect('/')

    # Tur başlangıç/bitiş tarihleri
    tur_baslangic_dt = secili_tur.get("BASLANGIC_TARIHI")   # datetime
    # datetime (zorunlu olmalı)
    tur_bitis_dt = secili_tur.get("BITIS_TARIHI")
    if not tur_baslangic_dt:
        flash("Tur başlangıç tarihi tanımlı değil.", "danger")
        return redirect('/')

    # Tarih stringleri (DD/MM/YYYY)
    tarih_sql_format = tur_baslangic_dt.strftime('%d/%m/%Y')
    bitis_sql_format = (tur_bitis_dt.strftime('%d/%m/%Y')
                        if tur_bitis_dt else tarih_sql_format)

    # Tur bilgisi: "TUR_ISMI BASLANGIC_TARIHI"
    tur_isim = (secili_tur.get("TUR_ISMI") or "").strip()
    tur_bilgisi = f"{tur_isim} {tarih_sql_format}".strip()

    tur_object_id = ObjectId(tur_id)

    # KVKK onay tarihi boşsa "şimdi"
    if not kvkk_onay_tarihi:
        kvkk_onay_tarihi = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

    # Kredi bilgisi
    kredi_bilgisi = db["kredi"].find_one({"TC_PASAPORT": tc_pasaport})
    kalan_kredi = kredi_bilgisi.get("KREDI", 0) if kredi_bilgisi else 0

    if kalan_kredi > 0:
        # Kredi düş
        db["kredi"].update_one({"TC_PASAPORT": tc_pasaport}, {
                               "$inc": {"KREDI": -1}})
        kalan_kredi -= 1

        # Rezervasyonu oluştur ve kaydet
        yeni_rezervasyon = _create_and_save_rezervasyon(
            form_data=request.form,
            tur=secili_tur,
            kredi_kullanildi=True,
            kalan_kredi_sonrasi=kalan_kredi
        )

        # E-postaları gönder
        _send_rezervasyon_emails(yeni_rezervasyon, secili_tur)

        flash(
            f'Kaydınız alınmıştır. Ödemeniz kredinizden düşülmüştür. Kalan Krediniz: {kalan_kredi}', 'success')
        return redirect('/')

    else:
        # Kredi yoksa ödeme akışına yönlendir
        return redirect(url_for(
            'odeme',
            adsoyad=adsoyad, telefon=telefon, email=email,
            tc_pasaport=tc_pasaport, dogum_tarihi=dogum_tarihi,
            adres=adres, sertifika=sertifika, tur_id=tur_id
        ))


# KAPASITE KONTROL ALANI


@app.route('/kontrol_kapasite', methods=['GET', 'POST'])
def kontrol_kapasite():
    # Hem JSON POST hem de GET query desteği
    if request.method == 'POST':
        data = request.get_json(silent=True) or {}
        tur_id = data.get('tur_id')
    else:
        tur_id = request.args.get('tur_id')

    if not tur_id:
        return jsonify({"error": "Gerekli parametre (tur_id) eksik."}), 400

    try:
        secili_tur = tur_collection.find_one({"_id": ObjectId(tur_id)})
        if not secili_tur:
            return jsonify({"error": "Tur bulunamadı"}), 404

        toplam_kontenjan = secili_tur.get("KONTENJAN", 0)
        kayitli_sayisi = customer_collection.count_documents(
            {"TUR_ID": ObjectId(tur_id)})
        musait_kontenjan = toplam_kontenjan - kayitli_sayisi

        return jsonify({
            "kayitli": kayitli_sayisi,
            "kapasite": toplam_kontenjan,
            "musait": max(0, musait_kontenjan)
        })
    except Exception as e:
        current_app.logger.error(f"Tur kapasite kontrol hatası: {e}")
        return jsonify({"error": "Tur kapasitesi kontrol edilirken bir hata oluştu."}), 500


@app.route('/kredi', methods=['GET', 'POST'])
def kredi_tanimla():
    if 'username' not in session:
        flash("Bu işlemi yapmak için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    if request.method == 'POST':
        tc_pasaport = request.form.get('tc_pasaport')
        ad = request.form.get('ad')
        email = request.form.get('email')
        kredi_baslama = request.form.get('kredi_baslama')
        kredi_bitis = request.form.get('kredi_bitis')
        is_update = (request.form.get('is_update') or '0').strip()

        try:
            girilen_kredi = int(request.form.get('kredi'))
        except (ValueError, TypeError):
            flash("Lütfen kredi için geçerli bir sayı girin.", "danger")
            return redirect(url_for('kredi_tanimla'))

        mevcut = db["kredi"].find_one({"TC_PASAPORT": tc_pasaport})

        if mevcut:
            if is_update == '1':
                # GÜNCELLE (SET): toplam kredi değeri girilen sayı olur
                yeni_kredi = girilen_kredi
            else:
                # EKLE (INCREMENT): mevcut üzerine ekle
                mevcut_kredi = int(mevcut.get("KREDI", 0) or 0)
                yeni_kredi = mevcut_kredi + girilen_kredi

            db["kredi"].update_one(
                {"TC_PASAPORT": tc_pasaport},
                {"$set": {
                    "AD": ad,
                    "EMAIL": email,
                    "KREDI": yeni_kredi,
                    "KREDI_BASLAMA": kredi_baslama,
                    "KREDI_BITIS": kredi_bitis
                }}
            )
            flash(f'Kredi güncellendi. Toplam Kredi: {yeni_kredi}', 'success')
        else:
            # İlk kayıt: doğrudan set
            db["kredi"].update_one(
                {"TC_PASAPORT": tc_pasaport},
                {"$set": {
                    "AD": ad,
                    "EMAIL": email,
                    "KREDI": girilen_kredi,
                    "KREDI_BASLAMA": kredi_baslama,
                    "KREDI_BITIS": kredi_bitis
                }},
                upsert=True
            )
            flash('Kredi başarıyla tanımlandı.', 'success')

        # Güncel listeyi görmek için aynı sayfaya dön
        return redirect(url_for('kredi_tanimla'))

    # Kredisi 0'dan büyük olan kullanıcıları listele
    kredi_listesi = list(db["kredi"].find({"KREDI": {"$gt": 0}}))
    return render_template('rez_kredi_vertical.html', kredi_listesi=kredi_listesi)


@app.route('/rapor_al', methods=['POST'])
def rapor_al():
    # 1. Yetkilendirme Kontrolü
    if 'username' not in session:
        flash("Bu işlemi yapmak için giriş yapmalısınız.", "warning")
        return redirect(url_for('login'))

    # 2. Veri Doğrulama
    try:
        yil = int(request.form.get('yil'))
        ay = request.form.get('ay')
    except (ValueError, TypeError):
        flash("Lütfen geçerli bir yıl ve ay seçin.", "danger")
        return redirect(url_for('rezervasyon_list'))

    # Filtreleme koşulları
    if ay == "Tum Aylar":
        # Yılın başlangıç ve bitiş tarihini oluştur
        start_date = datetime(yil, 1, 1)
        end_date = datetime(yil, 12, 31, 23, 59, 59)
        dosya_adi = f"gecmis_rezervasyon_raporu_{yil}.xlsx"
    else:
        ay = int(ay)
        start_date = datetime(yil, ay, 1)
        # Ayın son gününü dinamik olarak bul
        end_date = (start_date + pd.DateOffset(months=1)) - \
            pd.DateOffset(days=1)
        dosya_adi = f"gecmis_rezervasyon_raporu_{str(ay).zfill(2)}_{yil}.xlsx"
    filtre = {"TARIH_NESNESI": {"$gte": start_date, "$lte": end_date}}

    # MongoDB'den filtrelenmiş kayıtları al
    rapor_listesi = list(db["gecmis_rezervasyon"].find(filtre))

    # Verileri bir pandas DataFrame'e dönüştür
    data = []
    for rezervasyon in rapor_listesi:
        data.append({
            "Ad": rezervasyon.get("AD", ""),
            "Telefon": rezervasyon.get("GSM", ""),
            "E-posta": rezervasyon.get("EMAIL", ""),
            "TC/Pasaport No": rezervasyon.get("TC_PASAPORT", ""),
            "Doğum Tarihi": rezervasyon.get("DOGUM_TARIHI", ""),
            "Sertifika": rezervasyon.get("SERTIFIKA", ""),
            "Rezervasyon Tarihi": rezervasyon.get("TARIH", ""),
            "Kayıt Tarihi": rezervasyon.get("KAYIT_TARIHI", ""),
            "Tur Bilgisi": rezervasyon.get("TUR_BILGISI", ""),
            "Kredi": rezervasyon.get("KREDI", "")
        })

    # Eğer veri yoksa kullanıcıya bilgi ver
    if not data:
        flash("Seçilen tarih aralığında kayıt bulunamadı.", "warning")
        return redirect(url_for('rezervasyon_list'))

    df = pd.DataFrame(data)

    # "Sıra" sütununu ekle
    df.index += 1  # Sıra numaraları 1'den başlasın
    df.reset_index(inplace=True)
    df.rename(columns={"index": "Sıra"}, inplace=True)

    # Excel dosyasını oluştur
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Rapor")

    # Dosyayı kullanıcıya indirme olarak sun
    output.seek(0)
    response = make_response(output.read())
    response.headers["Content-Disposition"] = f"attachment; filename={dosya_adi}"
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return response


# Update the `/odeme_onayla` route to calculate `hash_degeri` correctly.
@app.route('/odeme_onayla', methods=['POST'])
def odeme_onayla():
    # Formdan gelen bilgileri al
    adsoyad = request.form.get('adsoyad')
    telefon = request.form.get('telefon')
    email = request.form.get('email')
    tc_pasaport = request.form.get('tc_pasaport')
    dogum_tarihi = request.form.get('dogum_tarihi')
    adres = request.form.get('adres')
    sertifika = request.form.get('sertifika')
    tur_id = request.form.get('tur_id')
    tarih = request.form.get('tarih')

    # Ödeme simülasyonu
    kart_no = request.form.get('kart_no')
    son_kullanma = request.form.get('son_kullanma')
    cvv = request.form.get('cvv')
    fiyat = request.form.get('odeme_tutari')

    if not all([kart_no, son_kullanma, cvv]):
        flash("Lütfen tüm ödeme bilgilerini giriniz.", "danger")
        return redirect(url_for("odeme"))

    tur = tur_collection.find_one(
        {"_id": ObjectId(tur_id)}) if tur_id else None
    if not tur:
        flash("Tur bilgisi bulunamadı.", "danger")
        return redirect('/')

    tarih_nesnesi = tur.get("BASLANGIC_TARIHI")
    tarih_sql_format = tarih_nesnesi.strftime('%d/%m/%Y')
    fiyat = tur.get("FIYAT", 0)  # Fetch the tour price

    # Rezervasyonu oluştur ve kaydet
    yeni_rezervasyon = _create_and_save_rezervasyon(
        form_data=request.form,
        tur=tur,
        kredi_kullanildi=False,
        kalan_kredi_sonrasi=0
    )

    # Loglama: ödeme detaylarını sakla
    odeme_log = {
        "TC_PASAPORT": tc_pasaport,
        "AD": adsoyad,
        "EMAIL": email,
        "TUR_ID": ObjectId(tur_id),
        "TUR_BILGISI": tur.get("TUR_ISMI"),
        "TARIH": tarih_sql_format,
        "ODEME_TARIHI": datetime.utcnow(),
        "ODEME_TUTARI": fiyat,  # Log the dynamic price
        "ODEME_TIPI": "Sanal POS",
        "ODEME_DURUMU": "Basarili",
        "IP_ADRESI": request.remote_addr,
        "ISLEM_ID": secrets.token_hex(12),
        "KART_SON_4_HANE": kart_no[-4:] if kart_no else "",
        "KAYNAK": "rezervasyon_formu",
        "ACIKLAMA": "Kredi Karti ödemesi ile rezervasyon yapıldı."
    }
    db["odeme_kayitlari"].insert_one(odeme_log)

    # E-postaları gönder
    _send_rezervasyon_emails(yeni_rezervasyon, tur)

    flash("Kaydınız ve ödemeniz alınmıştır.", "success")
    return redirect('/')


# Update the `/odeme` route to handle query parameters using `request.args`.
@app.route('/odeme')
def odeme():
    # GET parametrelerinden bilgileri al
    adsoyad = request.args.get('adsoyad')
    telefon = request.args.get('telefon')
    email = request.args.get('email')
    tc_pasaport = request.args.get('tc_pasaport')
    dogum_tarihi = request.args.get('dogum_tarihi')
    adres = request.args.get('adres')
    sertifika = request.args.get('sertifika')
    tur_id = request.args.get('tur_id')

    tur = tur_collection.find_one(
        {"_id": ObjectId(tur_id)}) if tur_id else None

    return render_template('odeme.html', adsoyad=adsoyad, telefon=telefon, email=email,
                           tc_pasaport=tc_pasaport, dogum_tarihi=dogum_tarihi, adres=adres,
                           sertifika=sertifika, tur=tur, fiyat=tur.get("FIYAT", 0))


if __name__ == '__main__':
    # pdf_yukle_ve_kaydet()#
    app.run(debug=True, host='10.251.3.11', port=5002)
